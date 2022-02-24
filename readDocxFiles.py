import SSS
import docx

# The program should be able to take a .docx document as input
# Then distinguish between tables and paragraph from the document
# This program should be able to read the contents of the document,
# and write them to another new document.

# When the new document is finished writing, the Google Drive API will upload it to the Google Drive.

# This is to test out how to perserve the structure of the input document.

# From working with this api, the paragraphs (docParagraphs) is an object.
# There should be a way to convert it to a string so that it can be written
# into the new document.

# Gets all the paragraphs in a document
def getParagraphs(document):
    docParagraphs = document.paragraphs
    return docParagraphs

# Gets all the tables in a document
def getTables(document):
    docTables = document.tables
    return docTables

# This will be the input document where the user would want to upload
doc = docx.Document("testDocuments/TestDocument.docx")

# paragraphs
p = getParagraphs(doc)

# tables
t = getTables(doc)

# The document where the input document will be copied
newDoc = docx.Document()

# Character in a document file
character = ''

# Stores the characters from the file
fileContent = ''

# Check that all paragraphs are extracted from document
for paragraph in p:
    print(paragraph.text)


# Need to see if its possible to check for a way to iterate through the file to get a character and have it encoded using SSS.

# Questions for when the document contains a table:
# Do we need to compute just the contents of the table or
# Is there a way to encrypt the entire table?
# What does an encrypted table look like?